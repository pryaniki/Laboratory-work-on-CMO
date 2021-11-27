from docx import Document
from docx.shared import Inches

from constants import set_constants as set_c
from solution import Application, Event, event_handler, get_data_for_an_calc
from get_data import get_conditions
from work_with_document import fill_table_for_report, fill_table_analysis_of_calculations


def write_report_on_task(n_task: int, document, tables: list[[list[Application] | list[Event]]], conditions: str):
    document.add_heading(f' Задание {n_task + 1}', 2)
    document.add_paragraph(conditions)
    dic = {
        1: 'Система массового обслуживания (D|M|n).',
        2: 'Система массового обслуживания (M|D|n).',
        3: 'Система массового обслуживания (M|M|n).'
    }
    print(tables[0])
    document.add_paragraph(dic[n_task+1])
    widths = (Inches(0.4), Inches(1), Inches(0.3), Inches(0.3), Inches(1), Inches(1), Inches(0.3))
    fill_table_for_report(document, tables[0], widths)
    document.add_paragraph(f'Таблица 2')
    widths = (Inches(0.4), Inches(1), Inches(0.3), Inches(0.3), Inches(1), Inches(1), Inches(0.3))
    fill_table_for_report(document, tables[1], widths)
    print(f'Выполнил задачу № {n_task + 1}')


def get_data_for_report() -> tuple:
    """  Получает данные для заполнения таблиц 1, 2 для задач 1, 2, 3, 4  """

    tables_task_1 = event_handler(1)
    tables_task_2 = event_handler(2)
    tables_task_3 = event_handler(3)

    tables_task_4 = [], [], []
    #tables_task_4 = get_data_for_an_calc([tables_task_1, tables_task_2, tables_task_3])

    return tables_task_1, tables_task_2, tables_task_3, tables_task_4


def create_report(variant, path_to_cond='lab_3.txt', doc_name='Report.doc'):
    """Заполняет черновую версию в файл doc_name"""
    data = get_conditions(variant, path_to_cond)
    print(data)
    variant = data.variant
    name = data.name
    set_c(*data.data)
    data_for_report = get_data_for_report()

    from constants import NUM_SMO, SERVICE_TIME, DELTA_T, LAMBD, MU

    conditions = f'Вариант №{variant}\n кол-во СМО = {NUM_SMO},T об={SERVICE_TIME}, Tз={DELTA_T} lambda={LAMBD}, mu = {MU}'
    print(conditions)

    document = Document()
    document.add_paragraph(name)
    document.add_paragraph(conditions)
    #write_report_on_task(0, document, data_for_report[0], conditions)
    #tables_task_2 = event_handler(2)
    #write_report_on_task(0, document, tables_task_2, conditions)
    for i in range(3):
        write_report_on_task(i, document, data_for_report[i], conditions)
    #fill_table_analysis_of_calculations(document, data_for_report[3])


    document.save(doc_name)
