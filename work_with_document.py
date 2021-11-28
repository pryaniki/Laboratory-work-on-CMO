"""
В модуле находятся функции для заполнения word документа

"""
import numpy as np

from ListWrapper import ListWrapper
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from constants import NUM_FOR_ROUND


def num_to_str(num: float | int, num_for_round=NUM_FOR_ROUND) -> str:
    """

    Преобразует число в строку округляя до NUM_FOR_ROUND знаков

    Если число -0.1**num_for_round < num < 0.1**num_for_round то возвращает '0'

    """

    if isinstance(num, str):
        return num
    else:
        eps = 0.1 ** num_for_round
        if -eps < num < eps:
            return '0'
        else:
            return str(np.round(num, num_for_round))


def _set_col_widths(table, widths):

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def set_size_fount(paragraph, size: int):
    """Задает размер шрифта для параграфа"""
    for run in paragraph.runs:
        font = run.font
        font.size = Pt(size)


def fill_table_for_report(document, data, widths):
    """Создает таблицу в документе и заполняет ее."""

    table = document.add_table(rows=len(data) + 1, cols=len(data[0]), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_col_widths(table, widths)
    for i, line in enumerate(data):
        for j, num in enumerate(line.get_data_for_report()):
            table.rows[i + 1].cells[j].text = num_to_str(num)

    return table


def list_to_list_wrapper(table):
    lst = []
    for el in table:
        lst.append(ListWrapper(el))
    return lst


def reverse_table(data: list):
    """ Переворачивает таблицу и добавляет по 1 строке в начало каждой строки таблицы"""

    res = []

    for i, row in enumerate(zip(*data)):
        lst = list(row)
        res.append(ListWrapper(lst))
    return res


def fill_table_analysis_of_calculations(document, data: list[list[ListWrapper]]):
    """
    Заполнить анализ расчетов


    """
    tables_3 = data[0]
    table_for_task_5 = data[2]
    frequency_table = data[3]
    vector_r = data[1]
    frequency_table_task_3 = data[4]

    document.add_heading(' Анализ результатов', 1)
    document.add_heading(' Система массового обслуживания (D|M|n)', 2)
    # t3
    dic = {
        1: 'Система массового обслуживания (D|M|n).',
        2: 'Система массового обслуживания (M|D|n).',
        3: 'Система массового обслуживания (M|M|n).'
    }
    widths = (Inches(2), Inches(1), Inches(1), Inches(1))
    for table, index in zip(tables_3, dic):

        document.add_paragraph(dic[index])
        _ = list_to_list_wrapper(table)
        fill_table_for_report(document, _, widths)
        if index != 3:
            _ = []
            for i, el in enumerate(frequency_table[index-1]):
                _.append(ListWrapper([i, el]))
            document.add_paragraph('Относительные частоты пребывания СМО в состояниях')

            fill_table_for_report(document, _, (Inches(1), Inches(1)))
    document.add_paragraph('задание 5')
    lst = list_to_list_wrapper(reverse_table(table_for_task_5))
    fill_table_for_report(document, lst, (Inches(1), Inches(1)))

    vec_r_text = ", ".join(num_to_str(float(num)) for num in vector_r)
    document.add_paragraph(f'r = ({vec_r_text})')

    lst = list_to_list_wrapper(reverse_table(reverse_table(frequency_table_task_3)))
    fill_table_for_report(document, lst, widths)
